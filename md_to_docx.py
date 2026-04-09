from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

MD_PATH = r"d:\04-Skill Grind\Projects\10 Fraud Detection Agent\PROJECT_DOCUMENT.md"
OUT_PATH = r"d:\04-Skill Grind\Projects\10 Fraud Detection Agent\PROJECT_DOCUMENT.docx"

doc = Document()

# ── Style helpers ────────────────────────────────────────────────────────────

def set_heading(doc, text, level):
    p = doc.add_heading(text.strip(), level=level)
    return p

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Inches(0.3)
    shading_elm = p._element
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def add_table_from_md(doc, rows):
    # Parse header and separator
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    data_rows = []
    for row in rows[2:]:  # skip separator line
        data_rows.append([c.strip() for c in row.strip('|').split('|')])

    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for r_idx, row in enumerate(data_rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            if c_idx < len(cells):
                cells[c_idx].text = val

def inline_run(paragraph, text):
    """Add text with **bold** and `code` inline formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

# ── Parse ────────────────────────────────────────────────────────────────────

with open(MD_PATH, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Headings
    if line.startswith('### '):
        set_heading(doc, line[4:], 3)
    elif line.startswith('## '):
        set_heading(doc, line[3:], 2)
    elif line.startswith('# '):
        set_heading(doc, line[2:], 1)

    # Code block
    elif line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_code_block(doc, code_lines)

    # Table
    elif line.strip().startswith('|'):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_rows.append(lines[i].rstrip('\n'))
            i += 1
        if len(table_rows) >= 3:
            add_table_from_md(doc, table_rows)
        continue

    # Horizontal rule
    elif re.match(r'^-{3,}$', line.strip()):
        doc.add_paragraph('─' * 60)

    # Bullet list
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        inline_run(p, line[2:])

    # Numbered list
    elif re.match(r'^\d+\. ', line):
        p = doc.add_paragraph(style='List Number')
        inline_run(p, re.sub(r'^\d+\. ', '', line))

    # Bold Q&A line (e.g. **Q1. ...**)
    elif line.strip().startswith('**') and line.strip().endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip()[2:-2])
        run.bold = True

    # Empty line
    elif line.strip() == '':
        pass  # natural paragraph break

    # Normal paragraph
    else:
        p = doc.add_paragraph()
        inline_run(p, line)

    i += 1

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
